#!/usr/bin/env python3
"""
DOCX Entity Extractor - Rule-Based Parser
Extracts financial entities from structured Word documents
"""

import re
import json
from docx import Document
from typing import Dict, List, Optional
from pathlib import Path


class DOCXEntityExtractor:
    """Rule-based entity extractor for financial DOCX documents"""

    # Entity field mappings (flexible matching)
    FIELD_MAPPINGS = {
        'Counterparty': ['party a'],
        'Initial Valuation Date': ['initial valuation date', 'initial value date'],
        'Notional': ['notional amount', 'notional', 'amount (n)', 'amount'],
        'Valuation Date': ['valuation date', 'final valuation date'],
        'Maturity': ['termination date', 'maturity', 'maturity date', 'final date'],
        'Underlying': ['underlying', 'reference asset', 'stock', 'security'],
        'Coupon': ['coupon', 'coupon (c)', 'interest rate', 'rate'],
        'Barrier': ['barrier', 'barrier (b)', 'barrier level'],
        'Calendar': ['business day', 'calendar', 'day convention']
    }

    def __init__(self, docx_path: str):
        """Initialize with path to DOCX file"""
        self.docx_path = Path(docx_path)
        self.doc = Document(str(self.docx_path))
        self.entities = {}

    def extract_all(self) -> Dict[str, any]:
        """Main extraction method - extracts all entities"""
        # Step 1: Extract from tables (primary source)
        self._extract_from_tables()

        # Step 2: Extract from document text (fallback)
        self._extract_from_text()

        # Step 3: Extract from document title (for some entities like Barrier)
        self._extract_from_title()

        # Step 4: Post-process and clean
        self._clean_entities()

        return self.entities

    def _extract_from_tables(self):
        """Extract entities from tables - primary method"""
        for table in self.doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]

                # Handle key-value pairs (2 columns)
                if len(cells) >= 2:
                    key = cells[0].strip().lower()
                    value = cells[1].strip()

                    # Skip empty values
                    if not value or value in ['', 'None', 'N/A']:
                        continue

                    # Match against entity types
                    for entity_name, field_variants in self.FIELD_MAPPINGS.items():
                        for variant in field_variants:
                            if variant in key:
                                # Store first match only
                                if entity_name not in self.entities:
                                    self.entities[entity_name] = value
                                break

    def _extract_from_text(self):
        """Extract entities from paragraph text (fallback)"""
        full_text = '\n'.join([para.text for para in self.doc.paragraphs])

        # Extract barrier percentage from title or text
        if 'Barrier' not in self.entities:
            barrier_match = re.search(r'[Bb]arrier\s+(\d+(?:\.\d+)?%)', full_text)
            if barrier_match:
                self.entities['Barrier'] = barrier_match.group(1)

    def _extract_from_title(self):
        """Extract entities from document title/header"""
        # First paragraph often contains title info
        if self.doc.paragraphs:
            title = self.doc.paragraphs[0].text

            # Extract barrier from title (e.g., "Barrier 75%")
            if 'Barrier' not in self.entities:
                barrier_match = re.search(r'[Bb]arrier\s+(\d+(?:\.\d+)?%)', title)
                if barrier_match:
                    self.entities['Barrier'] = barrier_match.group(1)

    def _clean_entities(self):
        """Clean and normalize extracted entities"""
        # Clean percentage values (ensure they have %)
        for key in ['Coupon', 'Barrier']:
            if key in self.entities:
                value = self.entities[key]
                # Add % if missing
                if '%' not in value and re.match(r'^\d+(?:\.\d+)?$', value):
                    self.entities[key] = value + '%'

        # Clean notional (standardize format)
        if 'Notional' in self.entities:
            self.entities['Notional'] = self.entities['Notional'].replace('(N)', '').strip()

    def to_json(self, indent=2) -> str:
        """Convert entities to JSON string"""
        return json.dumps(self.entities, indent=indent, ensure_ascii=False)

    def to_dict(self) -> Dict:
        """Return entities as dictionary"""
        return self.entities


def main():
    """Command-line interface"""
    import sys

    # Check arguments
    if len(sys.argv) < 2:
        print("Usage: python docx_parser.py <path_to_docx_file>")
        print("\nExample: python docx_parser.py ZF4894_ALV_07Aug2026_physical.docx")
        sys.exit(1)

    docx_file = sys.argv[1]

    # Check file exists
    if not Path(docx_file).exists():
        print(f"Error: File '{docx_file}' not found!")
        sys.exit(1)

    # Extract entities
    print(f"\n{'='*60}")
    print(f"DOCX Entity Extractor - Rule-Based Parser")
    print(f"{'='*60}")
    print(f"\nProcessing: {docx_file}\n")

    try:
        extractor = DOCXEntityExtractor(docx_file)
        entities = extractor.extract_all()

        # Display results
        print("Extracted Entities:")
        print("-" * 60)
        print(extractor.to_json())
        print("-" * 60)
        print(f"\nTotal entities extracted: {len(entities)}")

    except Exception as e:
        print(f"Error processing document: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
